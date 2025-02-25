import os
from typing import List, Dict
from tika import parser
import pdfplumber
from docx import Document
from utils.logger import Logger

class DocumentProcessor:
    def __init__(self, chunk_size: int = 512):
        self.logger = Logger.get_logger(__name__)
        self.chunk_size = chunk_size
        
    def parse_document(self, file_path: str) -> Dict:
        """解析不同格式的文档"""
        self.logger.info("解析文档: %s", file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                self.logger.debug("使用PDF解析器")
                return self._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                self.logger.debug("使用DOCX解析器")
                return self._parse_docx(file_path)
            elif ext == '.txt':
                self.logger.debug("使用TXT解析器")
                return self._parse_txt(file_path)
            else:
                self.logger.debug("使用Tika解析器")
                return self._parse_with_tika(file_path)
        except Exception as e:
            self.logger.error("文档处理错误: %s - %s", file_path, str(e))
            return {"content": "", "metadata": {}}

    def _parse_pdf(self, file_path: str) -> Dict:
        """解析PDF文件"""
        text = ""
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "pages": len(pdf.pages),
                "title": os.path.basename(file_path)
            }
            
            for page in pdf.pages:
                text += page.extract_text() 
                
        return {
            "content": text,
            "metadata": metadata
        }

    def _parse_docx(self, file_path: str) -> Dict:
        """解析DOCX文件"""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            "title": os.path.basename(file_path),
            "paragraphs": len(doc.paragraphs)
        }
        
        return {
            "content": text,
            "metadata": metadata
        }

    def _parse_txt(self, file_path: str) -> Dict:
        """解析TXT文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            
        metadata = {
            "title": os.path.basename(file_path)
        }
        
        return {
            "content": text,
            "metadata": metadata
        }

    def _parse_with_tika(self, file_path: str) -> Dict:
        """使用Tika解析其他格式文件"""
        parsed = parser.from_file(file_path)
        return {
            "content": parsed.get("content", ""),
            "metadata": parsed.get("metadata", {})
        }

    def create_chunks(self, text: str, overlap: int = 50) -> List[str]:
        """使用滑动窗口策略将文本分块"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # 如果不是最后一块，调整到最近的句子结束
            if end < len(text):
                # 查找下一个句子结束符
                next_period = text.find('.', end)
                if next_period != -1 and next_period - end < 100:  # 限制向后查找范围
                    end = next_period + 1
            
            chunk = text[start:end]
            chunks.append(chunk)
            
            # 移动滑动窗口，考虑重叠
            start = end - overlap

        return chunks 